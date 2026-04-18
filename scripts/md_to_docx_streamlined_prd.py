"""
One-shot converter: Ceasar_PRD_streamlined.md -> Ceasar_PRD_streamlined.docx

Hand-rolled markdown→docx converter focused on the structures actually
used in the streamlined PRD: H1/H2/H3 headings, paragraphs, bold/italic
inline runs, ordered + unordered lists, simple tables, code blocks,
blockquotes, and a hard-skip for the embedded mermaid block (Word
doesn't render mermaid; the standalone HTML companion is the visual
fallback). Designed to be re-runnable when the source PRD updates.

Usage:
    python -m scripts.md_to_docx_streamlined_prd

Reads:  planning/PRD/Ceasar_PRD_streamlined.md
Writes: planning/PRD/Ceasar_PRD_streamlined.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
SRC = REPO_ROOT / "planning" / "PRD" / "Ceasar_PRD_streamlined.md"
DST = REPO_ROOT / "planning" / "PRD" / "Ceasar_PRD_streamlined.docx"


# --- Inline run rendering -------------------------------------------------

# Match **bold**, *italic*, `code`, [text](url) — single-pass tokenizer.
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*"          # bold
    r"|\*[^*\s][^*]*\*"         # italic
    r"|`[^`]+`"                  # inline code
    r"|\[[^\]]+\]\([^)]+\))"     # link
)


def _add_inline_runs(paragraph, text: str) -> None:
    """Tokenize text into runs and apply inline styling."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
        elif part.startswith("[") and "](" in part:
            # [text](url) — render text with the URL appended in parens
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                paragraph.add_run(m.group(1))
                tail = paragraph.add_run(f" ({m.group(2)})")
                tail.italic = True
                tail.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


# --- Block-level conversion ----------------------------------------------


def convert(md_text: str) -> Document:
    doc = Document()

    # Tighter default style for body text
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page-ish header at the top
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SoundPulse — Streamlined Product Spec")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        "Plain-language fork of the master PRD — what SoundPulse is, "
        "how it makes money, and what runs today"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Note box about the diagram
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "📊 For the visual end-to-end lifecycle diagram, "
        "open Ceasar_PRD_streamlined_diagram.html in any browser "
        "(self-contained, no install needed)."
    )
    note_run.italic = True
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

    doc.add_paragraph()  # spacer

    lines = md_text.splitlines()
    i = 0
    skip_first_h1 = True  # md has its own title block; we replaced it above

    while i < len(lines):
        line = lines[i]

        # --- skip the YAML/blockquote header at the very top of the source
        if i == 0 and line.startswith("# "):
            i += 1
            continue
        if line.startswith("> ") and skip_first_h1:
            # Skip the leading blockquote intro
            while i < len(lines) and (lines[i].startswith("> ") or lines[i].strip() == ""):
                i += 1
            skip_first_h1 = False
            continue

        # --- skip horizontal rules
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # --- code fences (including the mermaid diagram — skip entirely)
        if line.startswith("```"):
            lang = line[3:].strip()
            i += 1
            block = []
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            if lang == "mermaid":
                # Mermaid doesn't render in Word; insert a placeholder note
                p = doc.add_paragraph()
                r = p.add_run(
                    "[End-to-end lifecycle diagram — see "
                    "Ceasar_PRD_streamlined_diagram.html for the visual "
                    "version, or render the mermaid block in the .md source.]"
                )
                r.italic = True
                r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            else:
                # Render plain code block as a monospaced paragraph
                p = doc.add_paragraph()
                cr = p.add_run("\n".join(block))
                cr.font.name = "Consolas"
                cr.font.size = Pt(9)
            continue

        # --- headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # --- table block
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1]):
            header_cells = [c.strip() for c in line.strip("|").split("|")]
            i += 2  # consume header + separator
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip("|").split("|")]
                table_rows.append(row)
                i += 1
            n_cols = len(header_cells)
            tbl = doc.add_table(rows=1, cols=n_cols)
            tbl.style = "Light Grid Accent 1"
            hdr_cells = tbl.rows[0].cells
            for c_idx, cell_text in enumerate(header_cells):
                run = hdr_cells[c_idx].paragraphs[0].add_run(cell_text)
                run.bold = True
            for row in table_rows:
                cells = tbl.add_row().cells
                for c_idx, cell_text in enumerate(row[:n_cols]):
                    _add_inline_runs(cells[c_idx].paragraphs[0], cell_text)
            doc.add_paragraph()  # spacer
            continue

        # --- bullet list
        if re.match(r"^\s*[-*]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                txt = re.sub(r"^\s*[-*]\s+", "", lines[i])
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, txt)
                i += 1
            continue

        # --- numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                txt = re.sub(r"^\s*\d+\.\s+", "", lines[i])
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, txt)
                i += 1
            continue

        # --- blockquote
        if line.startswith("> "):
            block = []
            while i < len(lines) and lines[i].startswith("> "):
                block.append(lines[i][2:])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(" ".join(block))
            run.italic = True
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            continue

        # --- blank line
        if not line.strip():
            i += 1
            continue

        # --- regular paragraph (may span multiple lines until blank)
        block = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not _looks_like_block_start(lines[i]):
            block.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(block))

    return doc


def _looks_like_block_start(line: str) -> bool:
    """Stop accumulating a paragraph when we hit something that starts a new block."""
    return (
        line.startswith("#")
        or line.startswith("```")
        or line.startswith("|")
        or line.startswith("> ")
        or re.match(r"^\s*[-*]\s+", line) is not None
        or re.match(r"^\s*\d+\.\s+", line) is not None
        or line.strip() in ("---", "***", "___")
    )


def main() -> int:
    md_text = SRC.read_text(encoding="utf-8")
    doc = convert(md_text)
    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"Wrote {DST}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
